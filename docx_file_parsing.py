import docx
import xml_elements
import xml_library
import xml_checkbox
import xml_multiple_choice
import xml_text_match
import xml_math_expression
import xml_numeric
import attributes
import QCWindow_shell as QC
import random
from enum import Enum

class question_type(Enum):
    SINGLE_CHOICE = 0
    MULTIPLE_CHOICE = 1
    NUMERIC = 2
    TEXT_MATCH = 3
    MATH_EXPRESSION = 4

class shuffle(Enum):
    NO_SHUFFLE = 5
    SHUFFLE = 6

class partial_credit(Enum):
    NO_PARTIAL_CREDIT = 7
    PARTIAL_CREDIT = 8

class IncorrectSyntax(Exception):
    pass

def MathJax_border_fix(question):
    i = 0
    while (i != -1 and i < len(question)):
        start_tag = question.find("$$", i)
        if start_tag != -1:
            end_tag = question.find("$$", start_tag + 2)
            if end_tag == -1:
                raise IncorrectSyntax("An odd number of \"$$\" in the question")
            if question.find('\n', start_tag, end_tag) !=-1:
                raise IncorrectSyntax("Line feed inside MathJax formula")
            question = question[: start_tag] + "\\(" + question[start_tag + 2: end_tag] + "\\)" + question[end_tag + 2 :]
        i = start_tag
    return question

def add_library_attributes(library, attr):
    library.add_property("url_name", attr.url_name)
    library.add_property("xblock-family", attr.xblock_family)
    library.add_property("video_auto_advance", attr.video_auto_advance)
    library.add_property("graded", attr.graded)
    library.add_property("tags", attr.tags)
    library.add_property("edxnotes", attr.edxnotes)
    library.add_property("use_latex_compiler", attr.use_latex_compiler)
    library.add_property("video_bumper", attr.video_bumper)
    library.add_property("course_edit_method", attr.course_edit_method)
    if attr.show_correctness != None:
        library.add_property("show_correctness", attr.show_correctness)
    library.add_property("static_asset_path", attr.static_asset_path)
    library.add_property("hide_from_toc", attr.hide_from_toc)
    library.add_property("in_entrance_exam", attr.in_entrance_exam)
    if attr.showanswer != None:
        library.add_property("showanswer", attr.showanswer)
    library.add_display_name(attr.display_name)
    library.add_property("group_access", attr.group_access)
    library.add_property("video_speed_optimizations", attr.video_speed_optimizations)
    library.add_property("graceperiod", attr.graceperiod)
    library.add_property("rerandomize", attr.rerandomize)
    library.add_property("user_partitions", attr.user_partitions)
    library.add_property("show_reset_button", attr.show_reset_button)
    library.add_property("days_early_for_beta", attr.days_early_for_beta)
    library.add_property("max_attempts", attr.max_attempts)
    library.add_property("self_paced", attr.self_paced)
    library.add_property("visible_to_staff_only", attr.visible_to_staff_only)
    library.add_property("org", attr.org)
    library.add_property("library", attr.library)
    library.add_inner_attributes()

def add_problem_attributes(problem, attr):
    problem.add_property("course_edit_method", attr.course_edit_method)
    problem.add_property("display_name", attr.display_name)
    problem.add_property("markdown", attr.markdown)
    problem.add_property("max_attempts", attr.max_attempts)
    problem.add_property("rerandomize", attr.rerandomize)
    problem.add_property("show_correctness", attr.show_correctness)
    problem.add_property("showanswer", attr.showanswer)
    problem.add_property("video_speed_optimizations", attr.video_speed_optimizations)
    problem.add_property("weight", attr.weight)
    if attr.submission_wait_seconds != None:
        problem.add_property("submission_wait_seconds", attr.submission_wait_seconds)

def write_specified_option(option_name, question_announcement):
    n = question_announcement.find(option_name)
    if n != -1 and (question_announcement.count('\"', 0, n) + question_announcement.count(chr(8220), 0, n) + question_announcement.count(chr(8221), 0, n)) % 2 == 0:
        n += len(option_name)
        while question_announcement[n] == " ":
            n += 1
        if question_announcement[n] == "=":
            n += 1
            while question_announcement[n] == " ":
                n += 1
            if question_announcement[n] == '\"' or ord(question_announcement[n]) == 8220 or ord(question_announcement[n]) == 8221:
                quote1 = question_announcement.find('\"', n+1)
                quote2 = question_announcement.find(chr(8220), n+1)
                quote3 = question_announcement.find(chr(8221), n+1)
                if quote1 == -1:
                    quote1 = len(question_announcement)+1
                if quote2 == -1:
                    quote2 = len(question_announcement)+1
                if quote3 == -1:
                    quote3 = len(question_announcement)+1
                end = min(quote1, quote2, quote3)
                if end != -1:
                    return MathJax_border_fix(question_announcement[n+1: end])


def create_checkbox_question(options, question):
    if question.count('\n*') == 0:
        raise IncorrectSyntax("Question does not have a correct answer")
    problem = xml_elements.problem()
    choiceresponse = xml_checkbox.choiceresponse(2, question)
    if options[2] == partial_credit.PARTIAL_CREDIT:
        choiceresponse.add_partial_credit()
    problem.add_content(choiceresponse)
    return problem

def create_multiple_choice_question(options, question):
    count = question.count('\n*')
    if count == 0:
        raise IncorrectSyntax("Question does not have a correct answer")
    if count > 1:
        raise IncorrectSyntax("Multiple correct answers are not allowed for this type of question")
    shuffle_set = False
    if options[1] == shuffle.SHUFFLE:
        shuffle_set = True
    problem = xml_elements.problem()
    multiple_choice_response = xml_multiple_choice.multiplechoiceresponse(2, question, shuffle_set)
    problem.add_content(multiple_choice_response)
    return problem

def create_text_match_question(question, attr):
    if question.count('\n*') == 0:
        raise IncorrectSyntax("Question does not have a correct answer")
    problem = xml_elements.problem()
    stringresponse = xml_text_match.stringresponse(2, question)
    if attr.reg_type != None:
        stringresponse.set_reg_type(attr.reg_type)
    if attr.trailing_text != None:
        stringresponse.set_trailing_text(attr.trailing_text)
    problem.add_content(stringresponse)
    return problem

def create_math_expression_question(question, attr):
    if question.count('\n*') == 0:
        raise IncorrectSyntax("Question does not have a correct answer")
    problem = xml_elements.problem()
    stringresponse = xml_math_expression.formularesponse(2, question)
    if attr.tolerance != None:
        stringresponse.set_tolerance(attr.tolerance)
    if attr.reg_type != None:
        stringresponse.set_reg_type(attr.reg_type)
    if attr.trailing_text != None:
        stringresponse.set_trailing_text(attr.trailing_text)
    problem.add_content(stringresponse)
    return problem

def create_numeric_question(question, attr):
    if question.count('\n*') == 0:
        raise IncorrectSyntax("Question does not have a correct answer")
    problem = xml_elements.problem()
    numericalresponse = xml_numeric.numericalresponse(2, question)
    if attr.tolerance != None:
        numericalresponse.set_tolerance(attr.tolerance)
    if attr.trailing_text != None:
        numericalresponse.set_trailing_text(attr.trailing_text)
    problem.add_content(numericalresponse)
    return problem

def random_problem_name():
    hex = ["0","1","2","3","4","5","6","7","8","9","a","b","c","d","e","f"]
    res = ""
    for i in range(32):
        res = res + random.choice(hex)
    return res

def parse_file(file_full_ref, library_attr, problems_attr, individually):
    doc = docx.Document(file_full_ref)

    #print(len(doc.paragraphs))

    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    united_text = str('\n'.join(text))
    #print(united_text)

    print("Conversion started")
    #Library creation
    lib = xml_library.library()
    p_names = []
    add_library_attributes(lib, library_attr)
    properties_order = ["display_name", "library", "org"]
    i = 0
    i_prop = 0
    #if i != united_text.find("Question 1"):
    #    if united_text[i + 1] != "\n":
    #        lib.add_property(properties_order[i_prop], united_text[i: united_text.find("\n")])
    #        i_prop += 1
    #        i = united_text.find("\n")
    #    else:
    #        i += 1
    #    while(i != united_text.find("\nQuestion 1")):
    #        #print(i)
    #        #print(united_text.find("\nQuestion 1"))
    #        if united_text[i + 1] != "\n":
    #            lib.add_property(properties_order[i_prop], united_text[i + 1: united_text.find("\n", i + 1)])
    #            i_prop += 1
    #            i = united_text.find("\n", i + 1)
    #        else:
    #            i += 1
    i = united_text.find("\nQuestion 1")
    if i == -1:
        i = united_text.find("Question 1")

    #Questions creation
    problems = list()
    q_i = 1
    try:
        while(i != -1 and i < len(united_text)):
            if united_text[i :i + 10 + len(str(q_i))] != "\nQuestion " + str(q_i):
                if i!=0 or united_text[i :i + 9 + len(str(q_i))] != "\nQuestion " + str(q_i):
                    raise IncorrectSyntax("Incorrect syntax of the question entry")
            options = [question_type.MULTIPLE_CHOICE, shuffle.NO_SHUFFLE, partial_credit.NO_PARTIAL_CREDIT]
            start_of_question = united_text.find("\n", i + 1)
            #start_of_answers = min(united_text.find("\n*"), united_text.find("\nA:"))
            end_of_question = united_text.find("\nQuestion", start_of_question)
            if end_of_question == -1:
                end_of_question = len(united_text)
            question_announcement = united_text[i + 1 : start_of_question]
            full_question = united_text[start_of_question : end_of_question]
            #print(question_announcement)
            #print(full_question)
            current_attr = attributes.refined_problem_attributes()

            #Type recognizing
            if question_announcement.find("multiple choice") != -1:
                if full_question.count('\n*') > 1:
                    options[0] = question_type.MULTIPLE_CHOICE
                else:
                    options[0] = question_type.SINGLE_CHOICE
            if question_announcement.find("single correct answer") != -1:
                options[0] = question_type.SINGLE_CHOICE
            if question_announcement.find("multiple correct answers") != -1:
                options[0] = question_type.MULTIPLE_CHOICE
            if question_announcement.find("checkbox") != -1:
                options[0] = question_type.MULTIPLE_CHOICE
            if question_announcement.find("text match") != -1:
                options[0] = question_type.TEXT_MATCH
            if question_announcement.find("math expression") != -1:
                options[0] = question_type.MATH_EXPRESSION
            if question_announcement.find("numeric") != -1:
                options[0] = question_type.NUMERIC
            if question_announcement.find("regex") != -1:
                raise IncorrectSyntax("Unsupported question type: regular expression")
            if question_announcement.find("regular expression") != -1:
                raise IncorrectSyntax("Unsupported question type: regular expression")
            if question_announcement.find("reflective single choice") != -1:
                raise IncorrectSyntax("Unsupported question type: reflective single choice")
            if question_announcement.find("reflective multiple choice") != -1:
                raise IncorrectSyntax("Unsupported question type: reflective multiple choice")
            if question_announcement.find("reflective text answer") != -1:
                raise IncorrectSyntax("Unsupported question type: reflective text answer")

            #Adding other options
            if question_announcement.find("shuffle") != -1:
                options[1] = shuffle.SHUFFLE
            if question_announcement.find("no shuffle") != -1:
                options[1] = shuffle.NO_SHUFFLE        
            if question_announcement.find("partial credit") != -1:
                options[2] = partial_credit.PARTIAL_CREDIT
            if question_announcement.find("no partial credit") != -1:
                options[2] = partial_credit.NO_PARTIAL_CREDIT

            full_full_question = question_announcement + full_question
            full_question = MathJax_border_fix(full_question)
            if options[0] ==  question_type.SINGLE_CHOICE:
                problems_attr.copy(current_attr)
                if individually:
                    QCWindow = QC.QCWindow_shell(current_attr, full_full_question)
                    QCWindow.exec_()
                    current_attr = QCWindow.get_attr()
                problem = create_multiple_choice_question(options, full_question)
            if options[0] ==  question_type.MULTIPLE_CHOICE:
                problems_attr.copy(current_attr)
                if individually:
                    QCWindow = QC.QCWindow_shell(current_attr, full_full_question)
                    QCWindow.exec_()
                    current_attr = QCWindow.get_attr()
                problem = create_checkbox_question(options, full_question)
            if options[0] ==  question_type.NUMERIC:
                problems_attr.copy(current_attr, "n")
                current_attr.tolerance = write_specified_option("tolerance", question_announcement)
                current_attr.trailing_text = write_specified_option("trailing_text", question_announcement)
                if individually:
                    QCWindow = QC.QCWindow_numeric(current_attr, full_full_question)
                    QCWindow.exec_()
                    current_attr = QCWindow.get_attr()
                problem = create_checkbox_question(options, full_question)
                problem = create_numeric_question(full_question, current_attr)
            if options[0] ==  question_type.MATH_EXPRESSION:
                problems_attr.copy(current_attr, "m")
                current_attr.tolerance = write_specified_option("tolerance", question_announcement)
                current_attr.reg_type = write_specified_option("type", question_announcement)
                current_attr.trailing_text = write_specified_option("trailing_text", question_announcement)
                if individually:
                    QCWindow = QC.QCWindow_math_expression(current_attr, full_full_question)
                    QCWindow.exec_()
                    current_attr = QCWindow.get_attr()
                problem = create_checkbox_question(options, full_question)
                problem = create_math_expression_question(full_question, current_attr)
            if options[0] ==  question_type.TEXT_MATCH:
                problems_attr.copy(current_attr, "t")
                current_attr.reg_type = write_specified_option("type", question_announcement)
                current_attr.trailing_text = write_specified_option("trailing_text", question_announcement)
                if individually:
                    QCWindow = QC.QCWindow_text_match(current_attr, full_full_question)
                    QCWindow.exec_()
                    current_attr = QCWindow.get_attr()
                problem = create_text_match_question(full_question, current_attr)

            add_problem_attributes(problem, current_attr)
            problem_text = []
            problem.add_to_text(problem_text)
            problem_str = '\n'.join(problem_text)
            #print(problem_str + '\n')
            p_name = random_problem_name()
            lib.add_problem(p_name)
            p_names.append(p_name)
            problems.append(problem_str)
            q_i += 1
            i = end_of_question


        lib_text = []
        lib.add_to_text(lib_text)
        lib_str = '\n'.join(lib_text)
        #print(lib_str + '\n')
        print("Conversion completed")
    except IncorrectSyntax as ex:
        print("The conversion was stopped for the following reasons:\n", ex.args[0], sep = "")
        lib_str = ""
        problems = []
    except xml_elements.IncorrectSyntax as ex:
        print("The conversion was stopped for the following reasons:\n", ex.args[0], sep = "")
        lib_str = ""
        problems = []

    return lib_str, problems, p_names

        